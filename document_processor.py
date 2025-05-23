"""
Document Processor Module
Handles document parsing, translation and rebuilding with format preservation.
"""
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import logging
import PyPDF2
import pdfplumber  # 新增pdfplumber库
from PIL import Image, ImageChops

# 添加PDF直接生成相关库
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# 添加docx2pdf导入
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    logging.warning("docx2pdf module not available. PDF conversion will be disabled.")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_docx(file_path):
    """
    Deep parse Word document structure and styles
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        dict: Document data structure containing styles, paragraphs, tables
    """
    logger.info(f"Parsing document: {file_path}")
    doc = Document(file_path)
    document_data = {
        'styles': {},
        'sections': [],
        'paragraphs': [],
        'tables': []
    }

    # Extract document global styles
    try:
        for style in doc.styles:
            if hasattr(style, 'type') and style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_data = {
                    'font_name': style.font.name if hasattr(style.font, 'name') else None,
                    'font_size': style.font.size.pt if hasattr(style.font, 'size') and style.font.size else None,
                    'bold': style.font.bold if hasattr(style.font, 'bold') else None,
                    'italic': style.font.italic if hasattr(style.font, 'italic') else None,
                    'color': str(style.font.color.rgb) if hasattr(style.font, 'color') and style.font.color and style.font.color.rgb else None,
                    'alignment': str(style.paragraph_format.alignment) if hasattr(style.paragraph_format, 'alignment') else None,
                    'line_spacing': style.paragraph_format.line_spacing if hasattr(style.paragraph_format, 'line_spacing') else None,
                }
                document_data['styles'][style.name] = style_data
    except Exception as e:
        logger.warning(f"Error extracting styles: {str(e)}")

    # Parse paragraphs and their styles
    for i, para in enumerate(doc.paragraphs):
        try:
            if not para.text.strip():  # Skip empty paragraphs
                continue
                
            para_data = {
                'text': para.text,
                'style': para.style.name if hasattr(para, 'style') else 'Normal',
                'runs': [],
                'format_markers': {
                    'is_title': para.style.name.lower().startswith('heading') 
                                or para.style.name.lower().startswith('title'),
                    'is_heading': para.style.name.lower().startswith('heading'),
                    'has_border': False,
                    'is_bold': False,
                    'alignment': 'left',
                    'font_size': 'normal',
                    'is_text_box': False,
                    'is_centered': para.alignment == 1 if hasattr(para, 'alignment') else False
                }
            }
            
            # Extract run-level formatting
            for run in para.runs:
                if not run.text.strip():  # Skip empty runs
                    continue
                    
                run_data = {
                    'text': run.text,
                    'font': run.font.name if hasattr(run.font, 'name') else None,
                    'size': run.font.size.pt if hasattr(run.font, 'size') and run.font.size else None,
                    'bold': run.font.bold if hasattr(run.font, 'bold') else None,
                    'italic': run.font.italic if hasattr(run.font, 'italic') else None,
                    'color': str(run.font.color.rgb) if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb else None,
                    'underline': run.font.underline if hasattr(run.font, 'underline') else None,
                }
                para_data['runs'].append(run_data)
                
            document_data['paragraphs'].append(para_data)
        except Exception as e:
            logger.warning(f"Error processing paragraph {i}: {str(e)}")

    # Parse table structures
    for i, table in enumerate(doc.tables):
        try:
            table_data = {
                'rows': [],
                'style': table.style.name if hasattr(table, 'style') and hasattr(table.style, 'name') else 'Table Grid'
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_data = {
                        'text': cell.text,
                        'paragraphs': [],
                        'style': cell.paragraphs[0].style.name if cell.paragraphs and hasattr(cell.paragraphs[0], 'style') else 'Normal'
                    }
                    
                    # Process cell paragraphs
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_data['paragraphs'].append(para.text)
                            
                    row_data.append(cell_data)
                table_data['rows'].append(row_data)
            document_data['tables'].append(table_data)
        except Exception as e:
            logger.warning(f"Error processing table {i}: {str(e)}")
    
    logger.info(f"Document parsing complete. Found {len(document_data['paragraphs'])} paragraphs and {len(document_data['tables'])} tables")
    return document_data

def extract_pdf_metadata(pdf_path):
    """提取PDF精确格式数据"""
    pdf_meta = {
        'pages': [],
        'fonts': set(),
        'global_margins': None,
        'text_blocks': []
    }
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            # 提取全局页边距（取第一页作为基准）
            first_page = pdf.pages[0]
            # 安全获取页边距
            try:
                pdf_meta['global_margins'] = {
                    'left': first_page.bbox[0] if hasattr(first_page, 'bbox') and first_page.bbox else 0,
                    'top': first_page.bbox[1] if hasattr(first_page, 'bbox') and first_page.bbox else 0,
                    'right': first_page.width - (first_page.bbox[2] if hasattr(first_page, 'bbox') and first_page.bbox else first_page.width),
                    'bottom': first_page.height - (first_page.bbox[3] if hasattr(first_page, 'bbox') and first_page.bbox else first_page.height)
                }
            except (AttributeError, IndexError):
                # 如果无法获取bbox，使用默认值
                pdf_meta['global_margins'] = {'left': 0, 'top': 0, 'right': 0, 'bottom': 0}
            
            # 逐页解析
            for page_num, page in enumerate(pdf.pages):
                page_data = {
                    'page_number': page_num + 1,
                    'size': (page.width, page.height),
                    'text_blocks': [],
                    'images': [],
                    'tables': [],
                    'headers_footers': []
                }
                
                # 文本块解析 - 改用更安全的方法
                try:
                    # 使用更稳健的方法从页面获取文本
                    # 方法一：尝试使用extract_words
                    try:
                        chars = []
                        for obj in page.chars:
                            if hasattr(obj, 'get'):
                                char_data = {
                                    'text': obj.get('text', ''),
                                    'x0': obj.get('x0', 0),
                                    'y0': obj.get('top', 0),
                                    'x1': obj.get('x1', 0),
                                    'bottom': obj.get('bottom', 0),
                                    'size': obj.get('size', 10)
                                }
                                chars.append(char_data)
                            else:
                                # 如果obj不是字典，尝试作为对象处理
                                try:
                                    char_data = {
                                        'text': getattr(obj, 'text', ''),
                                        'x0': getattr(obj, 'x0', 0),
                                        'y0': getattr(obj, 'top', 0),
                                        'x1': getattr(obj, 'x1', 0),
                                        'bottom': getattr(obj, 'bottom', 0),
                                        'size': getattr(obj, 'size', 10)
                                    }
                                    chars.append(char_data)
                                except:
                                    pass
                    except:
                        # 如果chars提取失败，尝试直接提取文本
                        chars = []
                    
                    # 如果上面的方法失败或没有字符，尝试使用extract_text
                    if not chars:
                        text_content = page.extract_text()
                        if text_content:
                            # 简单地将整个页面作为一个文本块
                            text_meta = {
                                'text': text_content,
                                'page': page_num + 1,
                                'bbox': [0, 0, page.width, page.height],
                                'format_markers': {
                                    'is_title': False,
                                    'is_text_box': False,
                                    'is_centered': False,
                                    'has_border': False
                                }
                            }
                            page_data['text_blocks'].append(text_meta)
                            pdf_meta['text_blocks'].append(text_meta)
                            
                            # 处理文本 - 分割成段落
                            paragraphs = text_content.split('\n\n')
                            for i, para in enumerate(paragraphs):
                                if not para.strip():
                                    continue
                                    
                                # 尝试检测特殊格式
                                is_title = False
                                is_text_box = False
                                is_centered = False
                                has_border = False
                                
                                # 检测标题 (通常在页面顶部，较短)
                                if i == 0 and len(para) < 100:
                                    is_title = True
                                
                                # 检测可能是特殊红框文本的内容
                                if "Biograph" in para or "Mission" in para:
                                    is_text_box = True
                                    has_border = True
                                    if len(para) < 100:  # 通常红框标题较短
                                        is_title = True
                                
                                # 创建文本块
                                para_meta = {
                                    'text': para,
                                    'page': page_num + 1,
                                    'bbox': [0, i*50, page.width, i*50+40],  # 估算位置
                                    'format_markers': {
                                        'is_title': is_title,
                                        'is_text_box': is_text_box,
                                        'is_centered': is_title,  # 标题通常居中
                                        'has_border': has_border
                                    }
                                }
                                
                                # 只有非空段落才添加
                                if para.strip():
                                    page_data['text_blocks'].append(para_meta)
                                    pdf_meta['text_blocks'].append(para_meta)
                    else:
                        # 将字符组合成单词
                        # 按y坐标(垂直位置)排序
                        lines = []
                        current_line = []
                        last_y = -1
                        
                        sorted_chars = sorted(chars, key=lambda x: (x['y0'], x['x0']))
                        
                        for char in sorted_chars:
                            # 如果是新的一行
                            if last_y == -1 or abs(char['y0'] - last_y) > 5:
                                if current_line:
                                    lines.append(current_line)
                                current_line = [char]
                                last_y = char['y0']
                            else:
                                current_line.append(char)
                                last_y = char['y0']
                        
                        # 添加最后一行
                        if current_line:
                            lines.append(current_line)
                        
                        # 将每行字符组合成文本块
                        for line_idx, line in enumerate(lines):
                            if not line:
                                continue
                                
                            # 排序行内字符（从左到右）
                            line.sort(key=lambda x: x['x0'])
                            
                            # 组合文本
                            line_text = ''.join([c['text'] for c in line if 'text' in c])
                            
                            # 计算行的边界框
                            x0 = min([c['x0'] for c in line if 'x0' in c], default=0)
                            y0 = min([c['y0'] for c in line if 'y0' in c], default=0)
                            x1 = max([c['x1'] for c in line if 'x1' in c], default=page.width)
                            y1 = max([c['bottom'] for c in line if 'bottom' in c], default=y0+10)
                            
                            # 检测格式（基于位置和内容）
                            is_title = False
                            is_text_box = False
                            is_centered = False
                            
                            # 检查是否居中
                            center_of_page = page.width / 2
                            line_center = (x0 + x1) / 2
                            if abs(line_center - center_of_page) < 100:  # 允许更宽的居中范围
                                is_centered = True
                            
                            # 检查是否为标题
                            if line_idx < 3 and len(line_text) < 100 and is_centered:
                                is_title = True
                            
                            # 检查是否为红框文本
                            if ("Biograph" in line_text or "Mission" in line_text or 
                                "WARNING" in line_text.upper() or "NOTICE" in line_text.upper()):
                                is_text_box = True
                            
                            # 创建文本块
                            text_meta = {
                                'text': line_text,
                                'page': page_num + 1,
                                'bbox': [x0, y0, x1, y1],
                                'format_markers': {
                                    'is_title': is_title,
                                    'is_text_box': is_text_box,
                                    'is_centered': is_centered,
                                    'has_border': is_text_box  # 假设文本框有边框
                                }
                            }
                            
                            # 只有非空文本才添加
                            if line_text.strip():
                                page_data['text_blocks'].append(text_meta)
                                pdf_meta['text_blocks'].append(text_meta)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num+1} using primary method: {str(e)}")
                    
                    # 退回到备用方法：直接提取文本
                    try:
                        text = page.extract_text()
                        if text:
                            # 简单分割成段落
                            paragraphs = text.split('\n\n')
                            y_offset = 0
                            for para in paragraphs:
                                if not para.strip():
                                    continue
                                
                                # 基本格式检测
                                is_title = False
                                is_text_box = False
                                
                                if "Biograph" in para or "Mission" in para:
                                    is_text_box = True
                                
                                if len(para) < 100 and y_offset < 100:
                                    is_title = True
                                
                                text_meta = {
                                    'text': para,
                                    'page': page_num + 1,
                                    'bbox': [50, y_offset, page.width-50, y_offset + 20],
                                    'format_markers': {
                                        'is_title': is_title,
                                        'is_text_box': is_text_box,
                                        'is_centered': is_title,
                                        'has_border': is_text_box
                                    }
                                }
                                page_data['text_blocks'].append(text_meta)
                                pdf_meta['text_blocks'].append(text_meta)
                                y_offset += 30  # 估算垂直间距
                    except Exception as backup_err:
                        logger.warning(f"Backup text extraction failed on page {page_num+1}: {str(backup_err)}")
                
                # 表格检测与提取 - 更安全的实现
                try:
                    tables = page.find_tables()
                    for table in tables:
                        # 安全检查表格属性
                        if not hasattr(table, 'rows') or not hasattr(table, 'cells'):
                            continue
                            
                        # 获取行列数
                        rows_count = len(table.rows) if hasattr(table, 'rows') else 0
                        cols_count = len(table.cols) if hasattr(table, 'cols') else 0
                        
                        # 如果无法确定列数，尝试从第一行推断
                        if cols_count == 0 and rows_count > 0:
                            if hasattr(table.rows[0], '__len__'):
                                cols_count = len(table.rows[0])
                            
                        table_data = {
                            'bbox': table.bbox if hasattr(table, 'bbox') else [0, 0, page.width, 100],
                            'rows': rows_count,
                            'cols': cols_count,
                            'cells': []
                        }
                        
                        # 只有当能获取单元格数据时才处理
                        if hasattr(table, 'cells') and table.cells:
                            for row_idx, row in enumerate(table.rows):
                                # 如果表格没有定义列属性，尝试从行获取列数
                                col_range = range(cols_count) if cols_count > 0 else (
                                    range(len(row)) if hasattr(row, '__len__') else range(0)
                                )
                                
                                for col_idx in col_range:
                                    try:
                                        # 安全获取单元格
                                        if row_idx < len(table.cells) and col_idx < len(table.cells[row_idx]):
                                            cell = table.cells[row_idx][col_idx]
                                            if cell is not None:
                                                # 提取文本
                                                try:
                                                    cell_text = cell.get_text().strip() if hasattr(cell, 'get_text') else str(cell)
                                                except:
                                                    cell_text = ""
                                                    
                                                if cell_text:
                                                    # 尝试获取单元格边界框
                                                    cell_bbox = getattr(cell, 'bbox', None)
                                                    if not cell_bbox:
                                                        # 估算位置
                                                        cell_bbox = [
                                                            col_idx * (page.width / cols_count),
                                                            row_idx * 20,
                                                            (col_idx + 1) * (page.width / cols_count),
                                                            row_idx * 20 + 20
                                                        ]
                                                        
                                                    cell_data = {
                                                        'row': row_idx,
                                                        'col': col_idx,
                                                        'text': cell_text,
                                                        'bbox': cell_bbox
                                                    }
                                                    table_data['cells'].append(cell_data)
                                    except Exception as cell_err:
                                        logger.debug(f"Error processing cell {row_idx}x{col_idx}: {str(cell_err)}")
                        
                        # 添加表格数据
                        if table_data['cells']:
                            page_data['tables'].append(table_data)
                except Exception as e:
                    logger.warning(f"Error extracting tables from page {page_num+1}: {str(e)}")
                
                # 添加页面数据
                pdf_meta['pages'].append(page_data)
        
        # 特殊格式检测 - 额外处理红框和标题
        for block in pdf_meta['text_blocks']:
            text = block['text']
            # 使用更精确的规则检测标题和红框文本
            if ("Mission Phase" in text or "Biograph" in text or 
                "PROCEDURE" in text.upper() or "PROTOCOL" in text.upper()):
                block['format_markers']['is_text_box'] = True
                block['format_markers']['has_border'] = True
                # 边框颜色设为红色
                block['format_markers']['border_color'] = 'red'
            
            # 检测可能是标题的内容
            if len(text.strip()) < 80 and any(keyword in text for keyword in ["Title:", "Chapter", "Section", "Mission"]):
                block['format_markers']['is_title'] = True
        
        logger.info(f"Advanced PDF extraction complete. Found {len(pdf_meta['pages'])} pages and {len(pdf_meta['text_blocks'])} text blocks")
        return pdf_meta
        
    except Exception as e:
        logger.error(f"Error in PDF metadata extraction: {str(e)}")
        # 返回空结构，避免完全失败
        return {
            'pages': [],
            'fonts': set(),
            'global_margins': {'left': 0, 'top': 0, 'right': 0, 'bottom': 0},
            'text_blocks': []
        }

def parse_pdf(file_path):
    """
    Parse PDF document content with enhanced formatting detection
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        dict: Document data structure containing text content with formatting information
    """
    logger.info(f"Parsing PDF document: {file_path}")
    
    # 使用新的提取方法
    pdf_meta = extract_pdf_metadata(file_path)
    
    # 转换为原有格式，保持兼容性
    document_data = {
        'pages': pdf_meta['pages'],
        'paragraphs': [],
        'tables': [],
        'text_boxes': [],
        'headings': []
    }
    
    # 将文本块转换为段落
    for block in pdf_meta['text_blocks']:
        para_data = {
            'text': block['text'],
            'page': block['page'],
            'bbox': block['bbox'],
            'format_markers': block['format_markers'],
            'runs': [{
                'text': block['text'],
                'bold': block['format_markers'].get('is_title', False),
                'italic': False
            }]
        }
        
        document_data['paragraphs'].append(para_data)
        
        # 特殊元素分类
        if block['format_markers'].get('is_title', False):
            document_data['headings'].append(para_data)
        elif block['format_markers'].get('is_text_box', False):
            document_data['text_boxes'].append(para_data)
    
    logger.info(f"Enhanced PDF parsing complete. Found {len(document_data['pages'])} pages, {len(document_data['paragraphs'])} paragraphs, {len(document_data['headings'])} headings, {len(document_data['text_boxes'])} text boxes")
    return document_data

def translate_document_content(document_data, translation_api_func, source_lang='zh', target_lang='en'):
    """
    Process document content through translation
    
    Args:
        document_data (dict): Parsed document data
        translation_api_func (callable): Function to call for translation
        source_lang (str): Source language code
        target_lang (str): Target language code
        
    Returns:
        dict: Document data with translated content
    """
    logger.info(f"Translating document from {source_lang} to {target_lang}")
    translated_doc = document_data.copy()
    translated_doc['source_lang'] = source_lang
    translated_doc['target_lang'] = target_lang
    
    # Translate paragraphs
    for i, para in enumerate(translated_doc['paragraphs']):
        try:
            if i % 10 == 0:
                logger.info(f"Translating paragraph {i+1}/{len(translated_doc['paragraphs'])}")
                
            original_text = para['text']
            translated_text = translation_api_func(original_text, source_lang, target_lang)
            para['translated_text'] = translated_text
            
            # Handle run-level translation if needed (for complex formatting)
            if len(para.get('runs', [])) > 1:
                current_position = 0
                for run in para['runs']:
                    # Calculate approximate position in translated text
                    run_length_ratio = len(run['text']) / len(original_text) if len(original_text) > 0 else 0
                    translated_length = int(len(translated_text) * run_length_ratio)
                    
                    # Handle boundary cases
                    end_position = min(current_position + translated_length, len(translated_text))
                    run['translated_text'] = translated_text[current_position:end_position]
                    current_position = end_position
                    
                    # Ensure the last run gets remaining text
                    if run == para['runs'][-1] and current_position < len(translated_text):
                        run['translated_text'] += translated_text[current_position:]
        except Exception as e:
            logger.error(f"Error translating paragraph {i}: {str(e)}")
            para['translated_text'] = para['text']  # Keep original on error
    
    # Translate tables
    for i, table in enumerate(translated_doc.get('tables', [])):
        try:
            logger.info(f"Translating table {i+1}/{len(translated_doc['tables'])}")
            for cell in table.get('cells', []):
                cell_text = cell.get('text', '')
                if cell_text.strip():
                    cell['translated_text'] = translation_api_func(cell_text, source_lang, target_lang)
        except Exception as e:
            logger.error(f"Error translating table {i}: {str(e)}")
    
    logger.info("Document translation complete")
    return translated_doc

def layout_compensation(original_text, translated_text):
    """文本长度差异补偿算法"""
    if not original_text or not translated_text:
        return {'scale': 1.0, 'tracking': 0}
        
    len_ratio = len(translated_text) / len(original_text)
    
    # 字体缩放策略
    if len_ratio > 1.2:
        return {'scale': 0.95, 'tracking': -0.5}
    elif len_ratio > 1.1:
        return {'scale': 0.98, 'tracking': -0.3}
    else:
        return {'scale': 1.0, 'tracking': 0}

def translate_pdf(input_path, output_path, translation_api_func, source_lang='zh', target_lang='en', percent_to_translate=100, output_pdf=False):
    """
    Complete translation workflow for PDF files with enhanced formatting
    
    Args:
        input_path (str): Path to source PDF document
        output_path (str): Path to save translated document (docx or pdf)
        translation_api_func (callable): Function to call for translation
        source_lang (str): Source language code
        target_lang (str): Target language code
        percent_to_translate (int): Percentage of document to translate (1-100)
        output_pdf (bool): Whether to output as PDF (requires docx2pdf)
        
    Returns:
        bool: Success status
    """
    try:
        # 1. Parse PDF with enhanced format detection
        pdf_meta = extract_pdf_metadata(input_path)
        pdf_meta['filepath'] = input_path
        
        # 转换为兼容格式
        doc_data = parse_pdf(input_path)
        
        # Apply percentage limit
        total_paragraphs = len(doc_data['paragraphs'])
        if percent_to_translate < 100:
            # Calculate how many paragraphs to translate
            paragraphs_to_translate = int(total_paragraphs * (percent_to_translate / 100))
            # Keep only the specified percentage
            doc_data['paragraphs'] = doc_data['paragraphs'][:paragraphs_to_translate]
            logger.info(f"Limiting translation to {percent_to_translate}% ({paragraphs_to_translate}/{total_paragraphs} paragraphs)")
        
        # 2. Translate content
        logger.info(f"Translating document from {source_lang} to {target_lang}")
        translated_data = translate_document_content(
            doc_data, 
            translation_api_func,
            source_lang, 
            target_lang
        )
        translated_data['percent'] = percent_to_translate
        translated_data['source_lang'] = source_lang
        translated_data['target_lang'] = target_lang
        
        # 3. Use the WordRebuilder to create a Word document
        # 确保输出路径使用.docx扩展名
        docx_output_path = output_path
        if output_pdf and output_path.lower().endswith('.pdf'):
            # 如果要输出PDF，创建临时DOCX文件
            docx_output_path = output_path.rsplit('.', 1)[0] + '.docx'
        
        rebuilder = WordRebuilder(pdf_meta, docx_output_path)
        docx_result_path = rebuilder.rebuild_document(translated_data)
        
        # 4. 如果需要PDF输出，转换DOCX为PDF
        if output_pdf:
            if not DOCX2PDF_AVAILABLE:
                logger.warning("PDF conversion requested but docx2pdf module is not available. Providing DOCX output only.")
                return docx_result_path is not None
            
            try:
                # 确保输出PDF路径
                pdf_output_path = output_path
                if not pdf_output_path.lower().endswith('.pdf'):
                    pdf_output_path += '.pdf'
                
                # 转换为PDF
                logger.info(f"Converting DOCX to PDF: {docx_result_path} -> {pdf_output_path}")
                convert(docx_result_path, pdf_output_path)
                
                # 如果转换成功且用户要求输出PDF，可以删除临时DOCX文件
                if os.path.exists(pdf_output_path) and docx_output_path != output_path:
                    os.remove(docx_output_path)
                    logger.info(f"Removed temporary DOCX file: {docx_output_path}")
                
                logger.info(f"PDF conversion complete: {pdf_output_path}")
                return os.path.exists(pdf_output_path)
            except Exception as pdf_err:
                logger.error(f"PDF conversion failed: {str(pdf_err)}")
                return docx_result_path is not None  # 返回DOCX结果状态
        
        return docx_result_path is not None
    except Exception as e:
        logger.error(f"PDF translation failed: {str(e)}")
        return False

def translate_docx(input_path, output_path, translation_api_func, source_lang='zh', target_lang='en', percent_to_translate=100):
    """
    Complete translation workflow for DOCX files
    
    Args:
        input_path (str): Path to source document
        output_path (str): Path to save translated document
        translation_api_func (callable): Function to call for translation
        source_lang (str): Source language code
        target_lang (str): Target language code
        percent_to_translate (int): Percentage of document to translate (1-100)
        
    Returns:
        bool: Success status
    """
    try:
        # 1. Parse document
        doc_data = parse_docx(input_path)
        
        # Apply percentage limit
        total_paragraphs = len(doc_data['paragraphs'])
        if percent_to_translate < 100:
            # Calculate how many paragraphs to translate
            paragraphs_to_translate = int(total_paragraphs * (percent_to_translate / 100))
            # Keep only the specified percentage
            doc_data['paragraphs'] = doc_data['paragraphs'][:paragraphs_to_translate]
            logger.info(f"Limiting translation to {percent_to_translate}% ({paragraphs_to_translate}/{total_paragraphs} paragraphs)")
        
        # 2. Translate content
        translated_data = translate_document_content(
            doc_data, 
            translation_api_func,
            source_lang, 
            target_lang
        )
        
        # 3. Rebuild document
        result_path = rebuild_document(translated_data, output_path, template_doc=input_path)
        
        # Add information about partial translation if needed
        if percent_to_translate < 100 and result_path:
            try:
                doc = Document(result_path)
                info_para = doc.add_paragraph()
                info_para.add_run(f"Note: This document contains {percent_to_translate}% of the original content.").bold = True
                info_para.add_run(f" Translated from {source_lang} to {target_lang}.")
                doc.save(result_path)
            except Exception as e:
                logger.warning(f"Could not add partial translation note: {str(e)}")
        
        return result_path is not None
    except Exception as e:
        logger.error(f"Document translation failed: {str(e)}")
        return False

def translate_document(input_path, output_path, translation_api_func, source_lang='zh', target_lang='en', percent_to_translate=100, direct_pdf=False):
    """
    Translate a document based on its file type
    
    Args:
        input_path (str): Path to source document
        output_path (str): Path to save translated document (docx or pdf)
        translation_api_func (callable): Function to call for translation
        source_lang (str): Source language code
        target_lang (str): Target language code
        percent_to_translate (int): Percentage of document to translate (1-100)
        direct_pdf (bool): Whether to directly generate PDF without DOCX
        
    Returns:
        bool: Success status
    """
    # Validate percentage value
    percent_to_translate = max(1, min(100, percent_to_translate))
    
    # Determine file type based on extension
    ext = os.path.splitext(input_path.lower())[1]
    
    if ext == '.docx':
        if direct_pdf and output_path.lower().endswith('.pdf'):
            return translate_docx_to_pdf_direct(input_path, output_path, translation_api_func, source_lang, target_lang, percent_to_translate)
        else:
            return translate_docx(input_path, output_path, translation_api_func, source_lang, target_lang, percent_to_translate)
    elif ext == '.pdf':
        if direct_pdf and output_path.lower().endswith('.pdf'):
            return translate_pdf_to_pdf_direct(input_path, output_path, translation_api_func, source_lang, target_lang, percent_to_translate)
        else:
            return translate_pdf(input_path, output_path, translation_api_func, source_lang, target_lang, percent_to_translate)
    else:
        logger.error(f"Unsupported file type: {ext}")
        return False

def translate_pdf_to_pdf_direct(input_path, output_path, translation_api_func, source_lang='zh', target_lang='en', percent_to_translate=100):
    """
    Translate PDF to PDF directly using reportlab
    
    Args:
        input_path (str): Path to source PDF document
        output_path (str): Path to save PDF document
        translation_api_func (callable): Function to call for translation
        source_lang (str): Source language code
        target_lang (str): Target language code
        percent_to_translate (int): Percentage of document to translate (1-100)
        
    Returns:
        bool: Success status
    """
    try:
        # 1. Parse PDF with enhanced format detection
        pdf_meta = extract_pdf_metadata(input_path)
        pdf_meta['filepath'] = input_path
        
        # 转换为兼容格式
        doc_data = parse_pdf(input_path)
        
        # Apply percentage limit
        total_paragraphs = len(doc_data['paragraphs'])
        if percent_to_translate < 100:
            # Calculate how many paragraphs to translate
            paragraphs_to_translate = int(total_paragraphs * (percent_to_translate / 100))
            # Keep only the specified percentage
            doc_data['paragraphs'] = doc_data['paragraphs'][:paragraphs_to_translate]
            logger.info(f"Limiting translation to {percent_to_translate}% ({paragraphs_to_translate}/{total_paragraphs} paragraphs)")
        
        # 2. Translate content
        logger.info(f"Translating document from {source_lang} to {target_lang}")
        translated_data = translate_document_content(
            doc_data, 
            translation_api_func,
            source_lang, 
            target_lang
        )
        translated_data['percent'] = percent_to_translate
        translated_data['source_lang'] = source_lang
        translated_data['target_lang'] = target_lang
        
        # 3. 直接使用reportlab生成PDF
        return generate_pdf_direct(translated_data, pdf_meta, output_path)
        
    except Exception as e:
        logger.error(f"PDF direct translation failed: {str(e)}")
        return False

def translate_docx_to_pdf_direct(input_path, output_path, translation_api_func, source_lang='zh', target_lang='en', percent_to_translate=100):
    """
    Translate DOCX to PDF directly using reportlab
    
    Args:
        input_path (str): Path to source DOCX document
        output_path (str): Path to save PDF document
        translation_api_func (callable): Function to call for translation
        source_lang (str): Source language code
        target_lang (str): Target language code
        percent_to_translate (int): Percentage of document to translate (1-100)
        
    Returns:
        bool: Success status
    """
    try:
        # 1. Parse document
        doc_data = parse_docx(input_path)
        
        # Apply percentage limit
        total_paragraphs = len(doc_data['paragraphs'])
        if percent_to_translate < 100:
            # Calculate how many paragraphs to translate
            paragraphs_to_translate = int(total_paragraphs * (percent_to_translate / 100))
            # Keep only the specified percentage
            doc_data['paragraphs'] = doc_data['paragraphs'][:paragraphs_to_translate]
            logger.info(f"Limiting translation to {percent_to_translate}% ({paragraphs_to_translate}/{total_paragraphs} paragraphs)")
        
        # 2. Translate content
        logger.info(f"Translating document from {source_lang} to {target_lang}")
        translated_data = translate_document_content(
            doc_data, 
            translation_api_func,
            source_lang, 
            target_lang
        )
        translated_data['percent'] = percent_to_translate
        translated_data['source_lang'] = source_lang
        translated_data['target_lang'] = target_lang
        
        # 3. 直接使用reportlab生成PDF
        # 创建一个简单的元数据结构，以便与PDF解析结果兼容
        simple_meta = {
            'filepath': input_path,
            'pages': [{'size': (612, 792)}],  # 默认Letter大小
            'global_margins': {'left': 36, 'top': 36, 'right': 36, 'bottom': 36}  # 默认边距0.5英寸
        }
        
        return generate_pdf_direct(translated_data, simple_meta, output_path)
        
    except Exception as e:
        logger.error(f"DOCX direct translation to PDF failed: {str(e)}")
        return False

def generate_pdf_direct(translated_data, pdf_meta, output_path):
    """
    Generate PDF document directly using reportlab
    
    Args:
        translated_data (dict): Translated document data
        pdf_meta (dict): PDF metadata with formatting information
        output_path (str): Path to save the PDF
        
    Returns:
        bool: Success status
    """
    try:
        logger.info(f"Generating PDF directly: {output_path}")
        
        # 获取页面大小信息
        page_width, page_height = (612, 792)  # 默认使用Letter大小
        first_page = pdf_meta.get('pages', [{}])[0]
        if first_page and 'size' in first_page and first_page['size']:
            # 使用PDF原始尺寸
            width, height = first_page['size']
            if width and height:
                page_width, page_height = width, height
        
        # 页面方向
        is_landscape = False
        if page_width > page_height:
            is_landscape = True
            page_width, page_height = page_height, page_width
        
        # 创建PDF文档
        doc = SimpleDocTemplate(
            output_path,
            pagesize=(page_width, page_height),
            rightMargin=pdf_meta.get('global_margins', {}).get('right', 36),
            leftMargin=pdf_meta.get('global_margins', {}).get('left', 36),
            topMargin=pdf_meta.get('global_margins', {}).get('top', 36),
            bottomMargin=pdf_meta.get('global_margins', {}).get('bottom', 36)
        )
        
        # 创建样式
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            alignment=TA_CENTER,
            fontSize=16,
            spaceAfter=12
        )
        normal_style = styles['Normal']
        info_style = ParagraphStyle(
            'Info',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.blue
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=16
        )
        text_box_style = ParagraphStyle(
            'TextBox',
            parent=styles['Normal'],
            backColor=colors.lightgrey,
            borderWidth=1,
            borderColor=colors.red,
            borderPadding=6,
            alignment=TA_CENTER,
            fontSize=12
        )
        
        # 创建内容列表
        content = []
        
        # 添加标题
        filename = os.path.basename(pdf_meta.get('filepath', 'document'))
        content.append(Paragraph(f"Translated Document: {filename}", title_style))
        
        # 添加信息区
        if 'source_lang' in translated_data:
            content.append(Paragraph(f"<b>Source language:</b> {translated_data['source_lang']}", info_style))
        if 'target_lang' in translated_data:
            content.append(Paragraph(f"<b>Target language:</b> {translated_data['target_lang']}", info_style))
        if 'percent' in translated_data and translated_data['percent'] < 100:
            content.append(Paragraph(f"<b>Translated content:</b> {translated_data['percent']}% of document", info_style))
        
        content.append(Spacer(1, 20))
        
        # 添加段落
        current_page = None
        for para in translated_data.get('paragraphs', []):
            if not para.get('translated_text'):
                continue
            
            # 如果有新的页码，添加页码标记
            if 'page' in para and para['page'] != current_page:
                current_page = para['page']
                content.append(Spacer(1, 10))
                content.append(Paragraph(f"<i>Page {current_page}</i>", info_style))
                content.append(Spacer(1, 5))
            
            # 根据格式标记使用不同的样式渲染
            format_markers = para.get('format_markers', {})
            if format_markers.get('is_title', False):
                # 标题样式
                content.append(Paragraph(para['translated_text'], heading_style))
            elif format_markers.get('is_text_box', False) or format_markers.get('has_border', False):
                # 带边框的文本框 - 检查是否为红框
                if format_markers.get('border_color') == 'red' or "Mission" in para['translated_text'] or "Biograph" in para['translated_text']:
                    # 使用表格模拟红框
                    data = [[Paragraph(para['translated_text'], text_box_style)]]
                    table = Table(data, colWidths=[doc.width - 12])
                    table.setStyle(TableStyle([
                        ('BOX', (0, 0), (-1, -1), 1, colors.red),
                        ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ]))
                    content.append(table)
                    content.append(Spacer(1, 10))
                else:
                    # 普通带边框文本
                    data = [[Paragraph(para['translated_text'], normal_style)]]
                    table = Table(data, colWidths=[doc.width - 12])
                    table.setStyle(TableStyle([
                        ('BOX', (0, 0), (-1, -1), 1, colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    content.append(table)
                    content.append(Spacer(1, 10))
            else:
                # 普通段落
                p_style = ParagraphStyle(
                    'CustomParagraph',
                    parent=normal_style,
                    alignment=TA_CENTER if format_markers.get('is_centered', False) else TA_LEFT
                )
                content.append(Paragraph(para['translated_text'], p_style))
                content.append(Spacer(1, 6))
        
        # 生成PDF
        doc.build(content)
        logger.info(f"Direct PDF generation complete: {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        return False

# PDF processing functions could be added here...

def get_format_preservation_prompt(text, source_lang, target_lang):
    """
    Generate a prompt that instructs the AI to preserve formatting markers
    
    Args:
        text (str): Text to translate with format markers
        source_lang (str): Source language
        target_lang (str): Target language
        
    Returns:
        str: Formatted prompt for translation with format preservation
    """
    return f"""Translate the following text from {source_lang} to {target_lang}, 
preserving all formatting markers exactly as they appear. 
Don't translate or modify anything inside format markers like <b>, <i>, etc.

TEXT TO TRANSLATE:
{text}

TRANSLATION (keep all formatting markers intact):
"""

def validate_format(pdf_meta, doc):
    """自动化格式验证"""
    report = {
        'text_blocks': {'passed': 0, 'failed': []},
        'page_layout': {'passed': 0, 'failed': []},
        'fonts': {'passed': 0, 'failed': []}
    }
    
    # 页面尺寸验证
    if hasattr(doc.sections[0], 'page_width') and pdf_meta.get('pages') and pdf_meta['pages'][0].get('size'):
        if abs(doc.sections[0].page_width - pdf_meta['pages'][0]['size'][0]) > 5:
            report['page_layout']['failed'].append('页面宽度偏差超过5pt')
    
    # 内容验证
    text_blocks_count = len(pdf_meta.get('text_blocks', []))
    paragraphs_count = len([p for p in doc.paragraphs if p.text.strip()])
    
    # 确认段落数量接近
    text_diff = abs(text_blocks_count - paragraphs_count)
    if text_blocks_count > 0 and text_diff / text_blocks_count > 0.2:  # 20%以上的差异
        report['text_blocks']['failed'].append(f'文本块数量差异过大: PDF={text_blocks_count}, DOCX={paragraphs_count}')
    else:
        report['text_blocks']['passed'] = min(text_blocks_count, paragraphs_count)
    
    # 页面布局验证
    if len(pdf_meta.get('pages', [])) != len(doc.sections):
        report['page_layout']['failed'].append(f'页数不匹配: PDF={len(pdf_meta.get("pages", []))}, DOCX={len(doc.sections)}')
    else:
        report['page_layout']['passed'] = len(pdf_meta.get('pages', []))
    
    return report

def visual_diff(pdf_image_path, docx_image_path, output_diff_path=None):
    """像素级差异检测"""
    try:
        img1 = Image.open(pdf_image_path).convert('RGB')
        img2 = Image.open(docx_image_path).convert('RGB')
        
        # 确保尺寸一致
        if img1.size != img2.size:
            img2 = img2.resize(img1.size)
        
        # 计算差异
        diff = ImageChops.difference(img1, img2)
        
        # 如果有差异且需要保存
        if diff.getbbox() and output_diff_path:
            diff.save(output_diff_path)
            return False, diff
        
        # 如果没有差异
        if not diff.getbbox():
            return True, None
            
        return False, diff
    except Exception as e:
        logger.error(f"Error in visual difference detection: {str(e)}")
        return False, None

def test_document_parser():
    """Simple test function to verify the document parser"""
    test_file = 'test_document.docx'
    if os.path.exists(test_file):
        data = parse_docx(test_file)
        print(f"Parsed {len(data['paragraphs'])} paragraphs and {len(data['tables'])} tables")
        return data
    else:
        print(f"Test file {test_file} not found")
        return None

class WordRebuilder:
    def __init__(self, pdf_meta, output_path):
        self.doc = Document()
        self.pdf_meta = pdf_meta
        self.output_path = output_path
        self._setup_document()
    
    def _setup_document(self):
        """设置基础文档结构"""
        # 获取第一页的尺寸
        first_page = self.pdf_meta['pages'][0] if self.pdf_meta['pages'] else None
        if not first_page:
            return
            
        # 页面方向
        if first_page.get('size') and first_page['size'][0] > first_page['size'][1]:
            section = self.doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
        
        # 页边距设置（如果有）
        if self.pdf_meta.get('global_margins'):
            section = self.doc.sections[0]
            margins = self.pdf_meta['global_margins']
            section.left_margin = Cm(margins['left']/28.35)  # 1cm=28.35pt
            section.right_margin = Cm(margins['right']/28.35)
            section.top_margin = Cm(margins['top']/28.35)
            section.bottom_margin = Cm(margins['bottom']/28.35)
            
        # 添加基本文档样式
        self._add_custom_styles()
    
    def _add_custom_styles(self):
        """添加自定义样式"""
        # 添加红框文本样式
        try:
            if 'RedBoxStyle' not in self.doc.styles:
                red_box_style = self.doc.styles.add_style('RedBoxStyle', WD_STYLE_TYPE.PARAGRAPH)
                red_box_style.font.bold = True
                red_box_style.font.size = Pt(12)
        except Exception as e:
            logger.warning(f"无法创建自定义样式: {str(e)}")
    
    def _create_red_box_table(self, text, page_num=None):
        """创建带红色边框的表格文本框"""
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        
        # 应用红色边框 (使用XML操作)
        try:
            # 设置表格边框属性
            tbl_pr = table._element.xpath('w:tblPr')[0]
            
            # 创建表格边框设置
            tbl_borders = OxmlElement('w:tblBorders')
            
            # 定义所有边框为红色
            for border_pos in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_pos}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # 边框宽度
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'FF0000')  # 红色边框
                tbl_borders.append(border)
            
            # 应用边框设置
            tbl_pr.append(tbl_borders)
            
            # 设置表格背景为浅灰色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')  # 浅灰色背景
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
            
        except Exception as e:
            logger.warning(f"无法应用红色边框样式: {str(e)}")
        
        # 页码引用
        if page_num:
            ref_para = cell.paragraphs[0]
            ref_para.add_run(f"[Page {page_num}] ").italic = True
            
        # 添加文本
        box_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = box_para.add_run(text)
        run.bold = True  # 加粗红框内文本
        
        # 居中对齐
        box_para.alignment = 1  # 居中
        
        # 增加段落间距
        self.doc.add_paragraph("")
        
        return table
    
    def rebuild_document(self, translated_data):
        """重建文档内容"""
        # 添加标题
        title_para = self.doc.add_heading(f"Translated PDF: {os.path.basename(self.pdf_meta.get('filepath', 'document'))}", level=1)
        title_para.alignment = 1  # 居中对齐
        
        # 添加信息区
        info_para = self.doc.add_paragraph()
        if 'source_lang' in translated_data:
            info_para.add_run(f"Source language: {translated_data['source_lang']}\n").bold = True
        if 'target_lang' in translated_data:
            info_para.add_run(f"Target language: {translated_data['target_lang']}\n").bold = True
        if 'percent' in translated_data:
            info_para.add_run(f"Translated content: {translated_data['percent']}% of document\n").bold = True
        
        self.doc.add_paragraph("=" * 50)
        
        # 如果没有段落，但有原始文本块，尝试使用它们创建简单文档
        if not translated_data.get('paragraphs') and self.pdf_meta.get('text_blocks'):
            for block in self.pdf_meta.get('text_blocks', []):
                text = block.get('text', '')
                if not text.strip():
                    continue
                
                # 根据格式标记使用不同的样式
                format_markers = block.get('format_markers', {})
                
                if format_markers.get('is_title', False):
                    # 标题格式
                    p = self.doc.add_heading(level=1)
                    if 'page' in block:
                        run = p.add_run(f"[Page {block['page']}] ")
                        run.italic = True
                    run = p.add_run(text)
                    run.bold = True
                    p.alignment = 1  # 居中
                
                elif format_markers.get('is_text_box', False) or format_markers.get('has_border', False):
                    # 判断是否需要红框
                    if format_markers.get('border_color') == 'red' or 'Mission' in text or 'Biograph' in text:
                        self._create_red_box_table(text, block.get('page'))
                    else:
                        # 普通带边框的文本框
                        table = self.doc.add_table(rows=1, cols=1)
                        table.style = 'Table Grid'
                        cell = table.cell(0, 0)
                        
                        # 页码引用
                        if 'page' in block:
                            ref_para = cell.paragraphs[0]
                            ref_para.add_run(f"[Page {block['page']}] ").italic = True
                            
                        # 添加文本
                        box_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        box_para.add_run(text)
                        
                        # 对齐方式
                        if format_markers.get('is_centered', False):
                            box_para.alignment = 1  # 居中
                        
                        # 段落间距
                        self.doc.add_paragraph("")
                
                else:
                    # 普通段落
                    p = self.doc.add_paragraph()
                    if 'page' in block:
                        p.add_run(f"[Page {block['page']}] ").italic = True
                        
                    p.add_run(text)
                    
                    # 应用对齐方式
                    if format_markers.get('is_centered', False):
                        p.alignment = 1  # 居中
        
        # 处理正常的翻译段落
        else:
            # 处理段落
            for para in translated_data.get('paragraphs', []):
                if not para.get('translated_text'):
                    continue
                    
                # 根据格式标记使用不同的样式
                if para.get('format_markers', {}).get('is_title', False):
                    # 标题格式
                    p = self.doc.add_heading(level=1)
                    if 'page' in para:
                        run = p.add_run(f"[Page {para['page']}] ")
                        run.italic = True
                    run = p.add_run(para['translated_text'])
                    run.bold = True
                    p.alignment = 1  # 居中
                    
                elif para.get('format_markers', {}).get('is_text_box', False) or para.get('format_markers', {}).get('has_border', False):
                    # 使用红框边框样式
                    if (para.get('format_markers', {}).get('border_color') == 'red' or 
                        "Mission" in para['translated_text'] or "Biograph" in para['translated_text']):
                        self._create_red_box_table(para['translated_text'], para.get('page'))
                    else:
                        # 普通文本框格式 - 使用单元格表格模拟
                        table = self.doc.add_table(rows=1, cols=1)
                        table.style = 'Table Grid'  # 添加边框
                        cell = table.cell(0, 0)
                        
                        # 页码引用
                        if 'page' in para:
                            ref_para = cell.paragraphs[0]
                            ref_para.add_run(f"[Page {para['page']}] ").italic = True
                            
                        # 翻译文本
                        box_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        box_para.add_run(para['translated_text'])
                        
                        # 对齐方式
                        if para.get('format_markers', {}).get('is_centered', False):
                            box_para.alignment = 1  # 居中
                        
                        # 段落间距
                        self.doc.add_paragraph("")
                    
                else:
                    # 普通段落
                    p = self.doc.add_paragraph()
                    if 'page' in para:
                        p.add_run(f"[Page {para['page']}] ").italic = True
                        
                    p.add_run(para['translated_text'])
                    
                    # 应用对齐方式
                    if para.get('format_markers', {}).get('is_centered', False):
                        p.alignment = 1  # 居中
        
        # 保存文档
        self.doc.save(self.output_path)
        logger.info(f"Enhanced document saved to: {self.output_path}")
        return self.output_path

def rebuild_document(translated_data, output_path, template_doc=None):
    """
    Rebuild document with translated content while preserving format
    
    Args:
        translated_data (dict): Translated document data
        output_path (str): Path to save the translated document
        template_doc (str, optional): Path to template document to use
        
    Returns:
        str: Path to the created document
    """
    logger.info(f"Rebuilding document to: {output_path}")
    
    # Create a new document or use template
    if template_doc:
        doc = Document(template_doc)
        # Clear content but keep styles
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
    else:
        doc = Document()
    
    # Add translated paragraphs
    for para_data in translated_data['paragraphs']:
        try:
            new_para = doc.add_paragraph()
            
            # Apply paragraph style if available
            if 'style' in para_data and para_data['style'] in doc.styles:
                new_para.style = doc.styles[para_data['style']]
            
            # If we have run-level data with formatting
            if 'runs' in para_data and para_data['runs']:
                for run_data in para_data['runs']:
                    if 'translated_text' in run_data and run_data['translated_text'].strip():
                        run = new_para.add_run(run_data['translated_text'])
                        
                        # 应用布局补偿（如果有）
                        if 'compensation' in para_data:
                            comp = para_data['compensation']
                            if comp.get('scale', 1.0) < 1.0:
                                if run_data.get('size'):
                                    scaled_size = run_data['size'] * comp['scale']
                                    run.font.size = Pt(scaled_size)
                        
                        # Apply run formatting
                        if run_data.get('font'):
                            run.font.name = run_data['font']
                        if run_data.get('size'):
                            run.font.size = Pt(run_data['size'])
                        if run_data.get('bold'):
                            run.font.bold = run_data['bold']
                        if run_data.get('italic'):
                            run.font.italic = run_data['italic']
                        if run_data.get('color'):
                            try:
                                # Convert RGB string to actual color
                                if isinstance(run_data['color'], str) and run_data['color'].startswith('RGB'):
                                    rgb_parts = run_data['color'].strip('RGB()').split(',')
                                    r, g, b = [int(x.strip()) for x in rgb_parts]
                                    run.font.color.rgb = RGBColor(r, g, b)
                            except Exception as e:
                                logger.warning(f"Could not apply color {run_data['color']}: {str(e)}")
            else:
                # Simple paragraph without complex formatting
                if 'translated_text' in para_data:
                    new_para.text = para_data['translated_text']
                    
            # 应用段落格式
            format_markers = para_data.get('format_markers', {})
            if format_markers.get('is_centered', False):
                new_para.alignment = 1  # 居中
        except Exception as e:
            logger.error(f"Error rebuilding paragraph: {str(e)}")
    
    # Add translated tables
    for table_data in translated_data.get('tables', []):
        try:
            if not table_data.get('rows', []):
                continue
                
            row_count = len(table_data['rows'])
            col_count = len(table_data['rows'][0]) if row_count > 0 else 0
            
            if row_count > 0 and col_count > 0:
                table = doc.add_table(rows=row_count, cols=col_count)
                
                # Apply table style if available
                if 'style' in table_data and table_data['style'] in doc.styles:
                    try:
                        table.style = table_data['style']
                    except Exception:
                        # Fallback to basic style if custom style fails
                        table.style = 'Table Grid'
                
                # Fill in cell contents
                for i, row in enumerate(table_data['rows']):
                    for j, cell_data in enumerate(row):
                        cell = table.cell(i, j)
                        
                        # Clear default paragraph
                        for p in cell.paragraphs:
                            p.clear()
                        
                        # Handle translated cell content
                        if 'translated_paragraphs' in cell_data and cell_data['translated_paragraphs']:
                            # Add each paragraph separately
                            for k, para_text in enumerate(cell_data['translated_paragraphs']):
                                if k == 0:
                                    # Use first paragraph
                                    cell.paragraphs[0].text = para_text
                                    # Apply style if available
                                    if 'style' in cell_data and cell_data['style'] in doc.styles:
                                        cell.paragraphs[0].style = doc.styles[cell_data['style']]
                                else:
                                    # Add additional paragraphs
                                    p = cell.add_paragraph(para_text)
                                    if 'style' in cell_data and cell_data['style'] in doc.styles:
                                        p.style = doc.styles[cell_data['style']]
                        elif 'translated_text' in cell_data:
                            # Single text for the whole cell
                            cell.text = cell_data['translated_text']
                            if 'style' in cell_data and cell_data['style'] in doc.styles:
                                cell.paragraphs[0].style = doc.styles[cell_data['style']]
        except Exception as e:
            logger.error(f"Error rebuilding table: {str(e)}")
    
    # Save the document
    try:
        doc.save(output_path)
        logger.info(f"Document saved to {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error saving document: {str(e)}")
        return None

if __name__ == "__main__":
    # When run directly, perform test parsing
    test_document_parser() 